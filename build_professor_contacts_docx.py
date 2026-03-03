import re
from io import BytesIO

import pdfplumber
import requests
from docx import Document

PDF_URL = (
    "https://education.gov.ng/wp-content/uploads/2022/12/"
    "2021-Directory-of-Full-Professors-in-the-Nigerian-University-System-FINAL.pdf"
)

RAW_OUTPUT_FILE = "professor_emails.txt"
CLEAN_EMAILS_FILE = "professor_emails_clean.txt"
DOCX_FILE = "professor_contacts_table.docx"

BAD_SUFFIXES = {
    "his",
    "her",
    "he",
    "dr",
    "prof",
    "mr",
    "mrs",
    "ms",
    "sir",
    "madam",
    "ibrahim",
    "onuh",
    "download",
    "directory",
    "docx",
    "doc",
}

STRICT_EMAIL_RE = re.compile(r"^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}$")
NAME_START_RE = re.compile(
    r"([A-Z][A-Z\-\.' ]+,\s*[A-Za-z][A-Za-z\-\.' ]+)\s+is a Professor"
)
EMAIL_AFTER_PHRASE_RE = re.compile(
    r"(?:His|Her)\s+email address is\s+([A-Za-z0-9._%+\-\s@]+(?:\.[A-Za-z]{2,}))",
    re.IGNORECASE | re.DOTALL,
)
GENERIC_EMAIL_RE = re.compile(
    r"[A-Za-z0-9._%+-]+\s*@\s*[A-Za-z0-9.-]+\s*\.\s*[A-Za-z]{2,}",
    re.IGNORECASE,
)


def normalize_spaces(text: str) -> str:
    return " ".join(text.split())


def normalize_name(name: str) -> str:
    return normalize_spaces(name).strip(" .,:;")


def normalize_email(raw_email: str) -> str:
    email = re.sub(r"\s+", "", raw_email.lower()).strip(".,;:()[]{}<>'\"")
    if "@" not in email:
        return ""

    local, domain = email.split("@", 1)
    while "." in domain:
        last_label = domain.rsplit(".", 1)[1]
        if last_label in BAD_SUFFIXES:
            domain = domain.rsplit(".", 1)[0]
        else:
            break
    return f"{local}@{domain}"


def load_clean_emails_from_output(path: str) -> list[str]:
    cleaned: list[str] = []
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        for line in f:
            raw = line.strip()
            if not raw:
                continue
            email = normalize_email(raw)
            if STRICT_EMAIL_RE.match(email):
                cleaned.append(email)

    seen: set[str] = set()
    deduped: list[str] = []
    for email in cleaned:
        if email not in seen:
            seen.add(email)
            deduped.append(email)
    return deduped


def extract_name_email_records(pdf_bytes: bytes) -> list[tuple[str, str]]:
    records: list[tuple[str, str]] = []
    with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
        full_text = "\n".join((p.extract_text() or "") for p in pdf.pages)

    name_matches = list(NAME_START_RE.finditer(full_text))
    for i, match in enumerate(name_matches):
        start = match.start()
        end = name_matches[i + 1].start() if i + 1 < len(name_matches) else len(full_text)
        block = full_text[start:end]
        name = normalize_name(match.group(1))

        email_candidate = ""
        m = EMAIL_AFTER_PHRASE_RE.search(block)
        if m:
            email_candidate = normalize_email(m.group(1))
        else:
            g = GENERIC_EMAIL_RE.search(block)
            if g:
                email_candidate = normalize_email(g.group(0))

        if STRICT_EMAIL_RE.match(email_candidate):
            records.append((name, email_candidate))

    deduped: dict[str, str] = {}
    for name, email in records:
        deduped.setdefault(email, name)

    return [(name, email) for email, name in deduped.items()]


def write_clean_email_file(records: list[tuple[str, str]]) -> None:
    emails = [email for _, email in records]
    with open(CLEAN_EMAILS_FILE, "w", encoding="utf-8") as f:
        f.write("\n".join(emails))


def write_docx_table(records: list[tuple[str, str]]) -> None:
    doc = Document()
    doc.add_heading("Directory of Full Professors (Extracted Contacts)", level=1)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Sn"
    header[1].text = "Name"
    header[2].text = "Email"

    for sn, (name, email) in enumerate(records, start=1):
        cells = table.add_row().cells
        cells[0].text = str(sn)
        cells[1].text = name if name else "N/A"
        cells[2].text = email

    doc.save(DOCX_FILE)


def main() -> None:
    print("Downloading PDF...")
    response = requests.get(PDF_URL, timeout=240)
    response.raise_for_status()

    print("Extracting name/email records...")
    parsed_records = extract_name_email_records(response.content)

    print(f"Loading and cleaning {RAW_OUTPUT_FILE}...")
    clean_emails = load_clean_emails_from_output(RAW_OUTPUT_FILE)

    name_by_email: dict[str, str] = {email: name for name, email in parsed_records}
    clean_set = set(clean_emails)

    records: list[tuple[str, str]] = []
    seen: set[str] = set()

    # Keep PDF order where name was parsed successfully.
    for name, email in parsed_records:
        if email in clean_set and email not in seen:
            records.append((name, email))
            seen.add(email)

    # Add remaining cleaned emails even if we could not parse the name.
    for email in clean_emails:
        if email not in seen:
            records.append((name_by_email.get(email, ""), email))
            seen.add(email)

    print(f"Writing {CLEAN_EMAILS_FILE}...")
    write_clean_email_file(records)

    print(f"Writing {DOCX_FILE}...")
    write_docx_table(records)

    print(f"Done. Saved {len(records)} records.")


if __name__ == "__main__":
    main()
